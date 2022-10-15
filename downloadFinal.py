

import os

import re

import numpy as np

import docx

from docx.shared import RGBColor
from docx import Document
from docx.shared import Pt

import similarity_algorithm

import UI.UserInterface


class download:
    global transre
    transre = None

    def text_create(name,flpaderpath):
        # Create final document
        rename = os.path.splitext(name)[0]
        desktop_path = flpaderpath ##report path
        full_path = desktop_path + "//" + rename + '_Report.txt'
        return full_path



    def data_matrix(fileroad,filename):
        # Open test file
        file = open(fileroad+'/'+filename,"r")
        row = file.readlines()
        print(row)
        l = []
        for line in row:
            str = list(line.rstrip())
            l.append(str)
        print(l)
        return l

    def text_write(filename, matrixname, searchmatri,ratio,filecheck):
        writedL=[]
        writedR=[]

        newRatio0 = float(ratio) #Keep two decimal places
        newRatio1 = newRatio0 * 100
        newRatio = float('%.2f'%newRatio1)

        row = len(matrixname) #Get the number of lines in the original file

        rows = np.array(searchmatri, dtype="object").shape[0]
        cols = np.array(searchmatri, dtype="object").shape[1]

        #Get the length of the longest line in the file
        mark = 0
        for i in range(0, row):
            Frames = matrixname[i]  # Set row I as the new array
            col = np.array(Frames).shape[0]  # Get new array length
            for j in range(0, col + 1):  # Loop through new array
                n = j
            if mark < n:  # Get longest row
                mark = n

        #Generate final report and mark duplicates
        with open(filename,'w') as f:
            f.write("-------------------------------------------------------------------")
            f.write("\n")
            f.write("This is the duplicate check report for the ")
            f.write(filecheck)
            f.write(" file")
            f.write("\n")
            f.write("The ratio is ")
            f.write(str(newRatio))
            f.write("%")
            f.write("\n")
            f.write("-------------------------------------------------------------------")
            f.write("\n")
            f.write("\n")

            writtenLine = 0
            for i in range(0,row+writtenLine): #Loop through the number of lines in the original file
                Frame = matrixname[i] #Set row I as the new array

                col = np.array(Frame).shape[0]  # Get new array length

                for j in range(0,col): #Loop through new array
                    Num = str(Frame[j])
                    f.write(Num)

                isWrite = 0

                #Duplicate tag
                for si in range(0,rows): # Get duplicate row range
                    newList = []
                    for sj in range(0,cols):
                        toList = list(searchmatri[si][sj])
                        newList.append(toList)
                    rowsCL = newList[0][0]
                    rowsCR = newList[1][0]

                    # rowsCR+=writtenLine
                    if rowsCL-1 == i: #Repeat first line
                        refile=""
                        for name in newList[2]:
                            refile+=name

                        # f.font.color.rgb = RGBColor(255,0,0)
                        if refile not in writedL:
                            if isWrite==0:
                                f.write("\n")
                                f.write(" #!# The next line is duplicated with: \"")
                                # for sj in range(colCL,col):
                                #     endNum = str(Frame[sj])
                                #     f.write(endNum)
                                f.write(refile)
                                f.write("\"")
                                isWrite = 1 #This row has been marked
                                writtenLine += 1
                            elif isWrite==1:
                                f.write(" and \"")
                                f.write(refile)
                                f.write("\"")
                            writedL.append(refile)

                    if (rowsCL-1 < i and i < rowsCR-1) and isWrite == 0 : #Repeat middle line

                        # f.write("\n")
                        # f.write("#!# The above line is duplicated with: \"")
                        # for refile in newList[2]:
                        #     f.write(refile)
                        # f.write("\"")
                        f.write(" #@# Repeated mark")
                        isWrite = 1
                    if i == rowsCR-1: #Repeat last line
                        refile = ""
                        for name in newList[2]:
                            refile += name

                        if refile not in writedR:
                            if isWrite==0:
                                f.write(" #@# Repeated mark")
                                f.write("\n")
                                f.write("#!# The above line is duplicated with: \"")
                                f.write(refile)
                                f.write("\"")
                                isWrite = 1
                                writtenLine += 1
                            elif isWrite==1:
                                f.write(" and \"")
                                f.write(refile)
                                f.write("\"")
                            writedR.append(refile)
                f.write("\n")
            f.close


    def mChange(exm1):
        #Modify the result format to facilitate marking
        i = 0
        NewList = []
        while i<len(exm1):
            ToList = list(exm1[i])
            NewList.append(ToList)
            i+=1
        return NewList

    def dictGet_key(exm1):
        return exm1.keys()

    def dictGet_value(exm1):
        return exm1.values()

    def changeDocx(textfile,desktop_path):
        with open(textfile,'r') as f:
            doc = Document()#new word
            p = doc.add_paragraph('')#Create a new paragraph, put this sentence outside the loop to reduce blank lines
            txtlines = f.readlines()


            for line in txtlines:
                if line.find("#!#")!=-1:
                    pt=r"(#!#)"#the split keywords and keep keywords
                    res = re.split(pt, line)#res[0] is character before keyword, res[1] is keywords, res[2] is character after keyword
                    run = p.add_run(res[1])
                    run.font.name=u'Calibri'
                    run.font.size = Pt(10)
                    r = run._element
                    run.font.color.rgb = RGBColor(250,0,0)

                    run = p.add_run(res[2])
                    run.font.name = u'Calibri'
                    run.font.size = Pt(10)
                    r = run._element
                    run.font.color.rgb = RGBColor(250, 0, 0)

                elif line.find("#@# Repeated mark")!=-1:
                    pt = "#@# Repeated mark"
                    res = re.split(pt, line)  # res[0] is character before keyword, res[1] is character after keyword
                    run = p.add_run(res[0])
                    run.font.name = u'Calibri'
                    run.font.size = Pt(10)
                    r = run._element
                    run.font.color.rgb = RGBColor(255, 69, 0)

                    run = p.add_run(res[1])
                    run.font.name = u'Calibri'
                    run.font.size = Pt(10)
                    r = run._element
                    run.font.color.rgb = RGBColor(255, 50, 0)
                else:
                    run = p.add_run(line)
                    run.font.name=u'Calibri'
                    run.font.size=Pt(10)
                    r = run._element
        rename = os.path.splitext(textfile)[0]
        doc.save(rename + '.docx')
        f.close()
        os.remove(textfile)



    def use(floader , names, needName, reportResult, flpaderpath):
        #Call method
        #download single files
        try:
            reportResults = reportResult
            allFlie = list(download.dictGet_key(reportResults[0]))
            allRate = list(download.dictGet_value(reportResults[0]))
            finalFile = list(download.dictGet_key(reportResults[1]))
            lines = list(download.dictGet_value(reportResults[1]))
            for i in range(0,len(allFlie)):
                if allFlie[i]==needName:
                    ReportFile = download.text_create(needName,flpaderpath)
                    fileName = str(names[i])
                    originalFile = download.data_matrix(floader, fileName)
                    repetitionRate = reportResults[0][allFlie[i]]
                    c1 = reportResults[1][allFlie[i]]
                    repetitionLine = download.mChange(c1)
                    download.text_write(ReportFile, originalFile, repetitionLine, repetitionRate, fileName)
                    download.changeDocx(ReportFile, flpaderpath)


        except Exception as e:
            print(e)

    def alluse(floader , names, reportResult, flpaderpath):
        #Call method
        #download all files
        try:
            reportResults = reportResult
            allFlie = list(download.dictGet_key(reportResults[0]))
            allRate = list(download.dictGet_value(reportResults[0]))
            finalFile = list(download.dictGet_key(reportResults[1]))
            lines = list(download.dictGet_value(reportResults[1]))
            for i in range(0,len(allFlie)):
                    ReportFile = download.text_create(names[i],flpaderpath)
                    fileName = str(allFlie[i])
                    originalFile = download.data_matrix(floader, fileName)
                    repetitionRate = reportResults[0][allFlie[i]]
                    c1 = reportResults[1][allFlie[i]]
                    repetitionLine = download.mChange(c1)
                    download.text_write(ReportFile, originalFile, repetitionLine, repetitionRate, fileName)
                    download.changeDocx(ReportFile,flpaderpath)

        except Exception as e:
            print(e)

    def trans (floader , names, needName, reportResult):
        #shown on the UI
        global transList
        transList=[]
        global transre
        try:
            transpath = desktop_path = os.path.join(os.path.expanduser('~'), "Desktop/")
            reportResults = reportResult
            allFlie = list(download.dictGet_key(reportResults[0]))
            allRate = list(download.dictGet_value(reportResults[0]))
            finalFile = list(download.dictGet_key(reportResults[1]))
            lines = list(download.dictGet_value(reportResults[1]))
            for i in range(0,len(allFlie)):
                if allFlie[i]==needName:
                    ReportFile = download.text_create(names[i], transpath)#file path
                    fileName = str(allFlie[i])#need file
                    originalFile = download.data_matrix(floader, fileName)#Detach the contents of the original file
                    repetitionRate = reportResults[0][allFlie[i]]#repeat rate
                    transre = reportResults[1][allFlie[i]]#get repeat lines
                    repetitionLine = download.mChange(transre)#repeat line
                    download.text_write(ReportFile, originalFile, repetitionLine, repetitionRate, fileName)
            with open(ReportFile, 'r') as f:
                for line in f:
                    line = line.strip('\n')
                    transList.append(line+'\n')
            f.close
            os.remove(ReportFile)
            return transList, transre



        except Exception as e:
            print(e)




